from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "docs/NX_Public_Progressive_Die_Workflow_Reorganized.docx"


SOURCES = [
    ("Siemens progressive die design overview",
     "https://www.siemens.com/en-gb/products/nx-manufacturing/tooling-fixture-design/progressive-die-design/"),
    ("Siemens Xcelerator Academy: Progressive Die Wizard Fundamentals",
     "https://training.plm.automation.siemens.com/mytraining/viewlibrary.cfm?product=XAPID298976455&track=103"),
    ("Siemens Xcelerator Academy: NX CAD Progressive Die Wizard Processes",
     "https://training.plm.automation.siemens.com/ilt/iltdescription.cfm?pID=TRCT2335______NX___10.0___2500"),
    ("NX Progressive Die Wizard Learning Hub index",
     "https://www.nxtop1.com/PDW/"),
    ("NX PDW workflow index",
     "https://www.nxtop1.com/PDW/workflow/index.html"),
    ("Demirezen Engineering strip-layout article",
     "https://www.demirezenengineering.com/blog/progressive-die-strip-layout"),
    ("Scientific Reports article on automatic strip layout design",
     "https://www.nature.com/articles/s41598-025-13328-1"),
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(55, 96, 146)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Public NX Progressive Die Workflow Reorganized for CadQuery Automation")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("An English working brief for Codex skills, CadQuery scripts, and CQ-Editor review")
r.font.name = "Arial"
r.font.size = Pt(10)
r.italic = True

doc.add_heading("Purpose", level=1)
doc.add_paragraph(
    "This document reorganizes publicly available Siemens/NX progressive-die information and related public engineering references into a practical, original workflow for AI-assisted CAD automation. "
    "It is not a copy of Siemens training material and it is not a replacement for NX Progressive Die Wizard, SolidWorks add-ons, AutoForm, or professional die-design review."
)

doc.add_heading("Public Sources Reviewed", level=1)
for name, url in SOURCES:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(name + ": ").bold = True
    p.add_run(url)

doc.add_heading("What the Public NX Sources Emphasize", level=1)
add_bullets(doc, [
    "Progressive die design is treated as a linked process, not just a set of independent solids.",
    "The workflow starts with sheet-metal part preparation and intermediate/staged part creation.",
    "Blank generation, blank layout, scrap design, and strip layout are central planning steps.",
    "Die tooling then grows from the strip: die base, standard parts, punches, inserts, openings, reliefs, and plates.",
    "Validation includes material usage, press force balance, clearances, reliefs, strip progression, motion, and collisions.",
    "Associativity matters: part updates should propagate through downstream die and process geometry where possible.",
    "Reusable company standards, templates, materials, part libraries, and tooling configurations are major productivity drivers.",
])

doc.add_heading("Reorganized Workflow", level=1)
workflow = [
    ("1. Part preparation", "Import target STEP, identify solids, measure bounding box/thickness, recognize holes/cutouts/forms, and choose a feed direction."),
    ("2. Blank generation", "Create or approximate the flat blank. For simple flat parts this may be the target outline; for formed parts it requires unfolding or forming analysis."),
    ("3. Blank layout", "Place one or more blanks on strip candidates and compare strip width, pitch, scrap bridge, edge web, nesting, and material utilization."),
    ("4. Scrap design", "Add carriers, scrap bridges, edge webs, slug areas, pilot holes, and separation logic so the strip remains feedable until final cutoff."),
    ("5. Strip layout", "Define stations along the feed direction and assign operations: pilot, pierce, trim, form, restrike, blank/cutoff, or idle."),
    ("6. Die base and standard parts", "Create upper/lower plates, guide pins, screws, dowels, springs, stripper plates, and standard components using templates or libraries."),
    ("7. Inserts and punches", "Generate piercing, blanking, trimming, and forming inserts from the strip geometry with clearance and relief assumptions."),
    ("8. Validation", "Check dimensions, clearances, material utilization, approximate force, center of pressure, collisions, slug paths, and station sequence."),
    ("9. Documentation and export", "Export STEP files, create drawings or reports, list assumptions, and record variant metrics."),
]
add_table(doc, ["Stage", "Open CadQuery/CQ-Editor interpretation"], workflow, [1.8, 5.8])

doc.add_heading("NX Concept to CadQuery Skill Mapping", level=1)
mapping = [
    ("Part preparation", "CadQuery importStep plus bounding-box, volume, face, and solid counts."),
    ("Blank generator", "Profile extraction or explicit 2D sketch/extrusion for simplified blanks."),
    ("Blank layout", "Parameterized strip pitch, strip width, and repeated transformed parts."),
    ("Scrap design", "Geometry for edge webs, scrap bridges, carriers, pilots, and slug zones."),
    ("Strip layout", "Station coordinate system and named CQ-Editor objects per station."),
    ("Die base management", "Simplified upper/lower plates and spacing parameters."),
    ("Standard parts", "Reusable CadQuery functions for guide pins, screws, dowels, and springs."),
    ("Piercing/forming inserts", "Extruded punches and lower die openings from cut profiles."),
    ("Force calculation", "Python formulas using perimeter, thickness, and shear strength."),
    ("Validation", "Re-imported STEP checks, bbox checks, clearance checks, and simple collision tests."),
]
add_table(doc, ["NX-style capability", "Open implementation approach"], mapping, [2.1, 5.5])

doc.add_heading("Inputs a Codex Skill Should Require", level=1)
add_bullets(doc, [
    "Target STEP file and whether it is the final part, a blank, or an assembly/component.",
    "Material name, sheet thickness, shear strength or conservative placeholder, and grain direction if relevant.",
    "Quantity or production volume, because this drives prototype tooling versus dedicated progressive tooling.",
    "Preferred process: blanking, piercing, bending, forming, drawing, compound die, transfer die, or progressive die.",
    "Strip assumptions: feed direction, station pitch, strip width, edge web, scrap bridge, carrier, and pilots.",
    "Tooling assumptions: clearance per side, plate thicknesses, punch height, die opening relief, guide pin layout, and open/shut height.",
    "Validation goals: dimension preservation, material use, center of pressure, force balance, collision avoidance, and manufacturability limits.",
])

doc.add_heading("Optimization Targets", level=1)
add_bullets(doc, [
    "Minimize stations while keeping operations physically feasible.",
    "Maximize material utilization by tuning strip pitch, strip width, blank orientation, and nesting.",
    "Balance forces around the press centerline or die center to reduce wear and accuracy problems.",
    "Maintain minimum edge web, scrap bridge, hole spacing, punch spacing, and die-wall strength.",
    "Avoid punch/die collisions and ensure slugs or formed features have relief paths.",
    "Keep generated tooling simple enough for real machining and assembly planning.",
])

doc.add_heading("Recommended GitHub Skill Structure", level=1)
add_numbered(doc, [
    "SKILL.md: trigger rules, scope, source policy, required inputs, geometry workflow, verification, and output conventions.",
    "references/input_checklist.md: target, material, production, die process, and expected outputs.",
    "references/progressive_die_workflow.md: public NX-style workflow rewritten for CadQuery.",
    "references/optimization_checks.md: material use, force, clearance, collision, and manufacturability checks.",
    "examples/: small CadQuery scripts showing one-step blanking and simplified progressive-die assemblies.",
])

doc.add_heading("Practical Boundary", level=1)
doc.add_paragraph(
    "CadQuery can generate geometry, CQ-Editor can display the models, and FreeCAD can inspect or translate STEP files. "
    "However, the manufacturing intelligence must come from explicit rules: company standards, public engineering formulas, professional judgment, or commercial systems such as NX Progressive Die Wizard and specialized die-design add-ons."
)

doc.add_heading("Immediate Next Steps for the CAD Repository", level=1)
add_bullets(doc, [
    "Keep the current progressive-die skill as the process-control layer.",
    "Add reusable CadQuery functions for strip, stations, pilots, punches, die blocks, guide pins, and assembly spacing.",
    "Add example scripts that export named STEP files and show the objects in CQ-Editor.",
    "Add a validation script that re-imports each STEP and reports bounding box, volume, and key assumptions.",
    "Add a short report template for every generated die concept.",
])

doc.save(OUT)
print(OUT)
