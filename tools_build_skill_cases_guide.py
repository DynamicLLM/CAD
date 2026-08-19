from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/CAD_Automation_Skill_Cases_Guide.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
        shade_cell(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


doc = Document()
s = doc.sections[0]
s.top_margin = Inches(0.8)
s.bottom_margin = Inches(0.8)
s.left_margin = Inches(0.85)
s.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.color.rgb = RGBColor(55, 96, 146)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CAD Automation Skill Cases Guide")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A modular skill structure for Codex, CadQuery, CQ-Editor, and manufacturing-process concepts")
r.font.name = "Arial"
r.font.size = Pt(10)
r.italic = True

doc.add_heading("Why Split The Skill Into Cases", level=1)
doc.add_paragraph(
    "A single broad progressive-die skill is useful for discussion, but practical CAD automation needs smaller cases. "
    "Flat blanking, bending, drawing/forming, fixturing, and machining all use geometry, but they require different assumptions, checks, and outputs. "
    "Case skills make Codex choose the right process before it writes CadQuery code."
)

doc.add_heading("Skill Library Overview", level=1)
rows = [
    ("cadquery-cad-foundation", "General CAD geometry, STEP import/export, measurement, CQ-Editor display.", "Any task that is not clearly a manufacturing-process case."),
    ("progressive-die-cadquery", "Parent workflow for NX-inspired progressive-die planning.", "When the user asks about progressive die overall."),
    ("progressive-die-flat-blanking", "Flat sheet products, piercing, trimming, pilot holes, final blanking.", "Best first case for 1.stp-style flat parts."),
    ("progressive-die-bending", "Bends, flanges, tabs, bend allowance, springback assumptions.", "When the part has simple formed angles."),
    ("progressive-die-drawing-forming", "Drawn cups, embosses, ribs, freeform forming.", "When forming physics may matter and must be marked conceptual."),
    ("fixture-design", "Locators, clamps, supports, datum pads, fixture assemblies.", "When making tooling to hold a target part."),
    ("machining-from-stock", "Stock block/bar, setup orientation, removal volume, fixture concept.", "When making a target shape from raw material by machining."),
    ("validation-and-optimization", "Variant checks, material use, force, collision, dimensions.", "After geometry exists or when optimizing parameters."),
]
table(doc, ["Skill", "Primary role", "When to use"], rows, [2.0, 2.8, 2.8])

doc.add_heading("Recommended Routing Logic", level=1)
bullets(doc, [
    "If the user only asks to create, inspect, convert, or display CAD geometry, use cadquery-cad-foundation.",
    "If the target is a flat sheet product with holes or perimeter cuts, use progressive-die-flat-blanking.",
    "If the target has simple bends or flanges, use progressive-die-bending.",
    "If the target has drawn, embossed, ribbed, or freeform sheet features, use progressive-die-drawing-forming.",
    "If the task is to hold a target part for machining, inspection, welding, or assembly, use fixture-design.",
    "If the task starts from a stock size and asks how to make a target shape, use machining-from-stock.",
    "If the user asks to compare variants, tune dimensions, or optimize material/process metrics, use validation-and-optimization.",
])

doc.add_heading("Expansion Plan By Case", level=1)
rows = [
    ("Foundation", "Reusable STEP import/export utilities, CQ-Editor viewer templates, measurement scripts."),
    ("Flat blanking", "Profile extraction from STEP, pilot placement, strip utilization, punch/die clearance offsets, blanking force."),
    ("Bending", "Bend allowance tables, K-factor presets, overbend parameters, staged bend displays."),
    ("Drawing/forming", "Intermediate-stage placeholders, draw-depth flags, forming envelope checks, external simulation handoff notes."),
    ("Fixture", "Datum selection helpers, clamp libraries, locator patterns, access/collision checks."),
    ("Machining", "Stock libraries, setup orientation, removal-volume estimate, fixture/vise templates."),
    ("Validation", "Variant table generation, constraints file, simple collision and force checks, report templates."),
]
table(doc, ["Case", "Next useful expansion"], rows, [1.8, 5.8])

doc.add_heading("Example Workflow For 1.stp", level=1)
doc.add_paragraph(
    "The current 1.stp example should use the flat-blanking case first. The workflow is to import the target, derive or approximate its flat profiles, create raw strip material, create piercing punches for internal holes or cutouts, create a lower die opening with clearance, create a final blanking punch for the outer profile, export STEP files, and inspect the assembly in CQ-Editor."
)

doc.add_heading("Boundary And Safety", level=1)
bullets(doc, [
    "These skills generate conceptual geometry and process logic; they are not certified tooling-design software.",
    "Public NX documents should be used as workflow inspiration only, not copied into the repository.",
    "For production die design, an engineer must confirm material data, press capacity, clearances, station layout, and tooling standards.",
    "CadQuery is strongest for parameterized geometry; CQ-Editor is strongest for quick visual review; FreeCAD remains useful for external STEP inspection and translation.",
])

doc.add_heading("GitHub Organization", level=1)
doc.add_paragraph("The repository now uses this skill layout:")
for line in [
    "skills/cadquery-cad-foundation/",
    "skills/progressive-die-cadquery/",
    "skills/progressive-die-flat-blanking/",
    "skills/progressive-die-bending/",
    "skills/progressive-die-drawing-forming/",
    "skills/fixture-design/",
    "skills/machining-from-stock/",
    "skills/validation-and-optimization/",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.save(OUT)
print(OUT)
